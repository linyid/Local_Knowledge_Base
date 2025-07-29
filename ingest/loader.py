# Contains libraries/functions that handle per-type logic (can later be extended modularly)
from pathlib import Path
from typing import Optional
from ingest.document import Document
import fitz  # PyMuPDF
import docx

def load_file(path: Path) -> Optional[Document]:
    """
    Loads a file from the given path and returns a Document object.

    Args:
        path (Path): The path to the file to be loaded.
    Returns:
        Optional[Document]: A Document object containing the file's content and metadata, or None if loading fails.
    """
    try:
        if path.suffix.lower() == '.txt':
            return load_txt(path)
        elif path.suffix.lower() == '.pdf':
            return load_pdf(path)
        elif path.suffix.lower() == '.docx':
            return load_docx(path)
        else:
            print(f"Unsupported file type: {path.suffix}")
            return None
    except Exception as e:
        print(f"Error loading file {path}: {e}")
    return None
    
def load_txt(path: Path) -> Document:
    """
    Loads a text file and returns a Document object.

    Args:
        path (Path): The path to the text file.
    Returns:
        Document: A Document object containing the file's content and metadata.
    """
    text = path.read_text(encoding='utf-8', errors='ignore')
    return Document(
        content=text,
        source_path=path,
        file_type='txt'
    )

def load_pdf(path: Path) -> Document:
    """
    Loads a PDF file and returns a Document object.

    Args:
        path (Path): The path to the PDF file.
    Returns:
        Document: A Document object containing the file's content and metadata.
    """
    doc = fitz.open(path)
    content = ""
    for page in doc:
        content += page.get_text()
    doc.close()
    return Document(
        content=content,
        source_path=path,
        file_type='pdf'
    )

def load_docx(path: Path) -> Document:
    """
    Loads a DOCX file and returns a Document object.

    Args:
        path (Path): The path to the DOCX file.
    Returns:
        Document: A Document object containing the file's content and metadata.
    """
    doc = docx.Document(path)
    content = "\n".join(para.text for para in doc.paragraphs)
    return Document(
        content=content,
        source_path=path,
        file_type='docx'
    )